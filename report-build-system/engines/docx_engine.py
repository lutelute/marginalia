"""
python-docx based build engine.

Generates high-quality DOCX output by directly manipulating Word XML,
supporting SEQ field auto-numbering, OMML equations, and template injection.
"""

import re
from pathlib import Path
from typing import Any, Dict, List, Optional

from . import BuildEngine, register_engine

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from lxml import etree

    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


@register_engine('python-docx')
class DocxEngine(BuildEngine):
    """Build DOCX documents using python-docx for direct Word XML manipulation."""

    def __init__(self):
        self._fig_counter = 0
        self._tbl_counter = 0
        self._eq_counter = 0
        self._alg_counter = 0
        self._labels: Dict[str, str] = {}  # label -> "Figure 5-1" etc.

    # -----------------------------------------------------------------
    # BuildEngine interface
    # -----------------------------------------------------------------

    def build(
        self,
        manifest: Dict[str, Any],
        format_type: str,
        content: str,
        manifest_path: Path,
        project_root: Path,
    ) -> bool:
        if format_type != 'docx':
            print(f"ERROR: python-docx engine only supports DOCX output, got '{format_type}'")
            return False

        if not HAS_PYTHON_DOCX:
            print("ERROR: python-docx is not installed. Run: pip install python-docx lxml")
            return False

        # Reset counters
        self._fig_counter = 0
        self._tbl_counter = 0
        self._eq_counter = 0
        self._alg_counter = 0
        self._labels = {}

        output_dir = project_root / 'output'
        output_dir.mkdir(exist_ok=True)
        manifest_name = manifest_path.stem
        output_path = output_dir / f"{manifest_name}.docx"

        # Resolve template
        template_path = self._resolve_inject_template(manifest, project_root)

        # Load or create document
        if template_path and template_path.exists():
            doc = Document(str(template_path))
        else:
            doc = Document()

        # Configuration from manifest
        docx_direct = manifest.get('docx-direct', {})
        anchor_heading = docx_direct.get('anchor-heading')
        chapter_prefix = docx_direct.get('chapter-prefix')
        crossref_mode = docx_direct.get('crossref-mode', 'seq')
        first_line_indent = docx_direct.get('first-line-indent', 0)
        page_break_before_h2 = docx_direct.get('page-break-before-h2', False)

        # Strip YAML front matter from content
        content = self._strip_front_matter(content)

        # Find injection point
        insert_idx = self._find_anchor(doc, anchor_heading)

        rel_output = output_path.relative_to(project_root)
        print(f"Building {rel_output} (python-docx)... ", end='', flush=True)

        try:
            self._render_markdown(
                doc=doc,
                content=content,
                insert_idx=insert_idx,
                chapter_prefix=chapter_prefix,
                crossref_mode=crossref_mode,
                first_line_indent=first_line_indent,
                page_break_before_h2=page_break_before_h2,
                project_root=project_root,
            )

            doc.save(str(output_path))
            print("\u2713")
            return True

        except Exception as e:
            print("FAILED")
            print(f"ERROR: {e}")
            return False

    def check_dependencies(self) -> Dict[str, bool]:
        return {
            'python-docx': HAS_PYTHON_DOCX,
            'lxml': HAS_PYTHON_DOCX,  # lxml is required by python-docx
        }

    def resolve_template(
        self,
        manifest: Dict[str, Any],
        format_type: str,
        project_root: Path,
    ) -> Optional[Path]:
        return self._resolve_inject_template(manifest, project_root)

    # -----------------------------------------------------------------
    # Template resolution
    # -----------------------------------------------------------------

    def _resolve_inject_template(
        self, manifest: Dict[str, Any], project_root: Path,
    ) -> Optional[Path]:
        """Resolve python-docx inject template path."""
        # Check bundle info
        bundle = manifest.get('_bundle')
        if bundle and 'python-docx' in bundle:
            docx_bundle = bundle['python-docx']
            f = docx_bundle.get('docx')
            if f:
                p = project_root / 'templates' / f
                if p.exists():
                    return p

        # Default convention: templates/docx/{template}-inject.docx
        template = manifest.get('template', 'report')
        style = manifest.get('style')
        if style:
            styled = project_root / 'templates' / 'docx' / f"{template}-{style}-inject.docx"
            if styled.exists():
                return styled
        base = project_root / 'templates' / 'docx' / f"{template}-inject.docx"
        if base.exists():
            return base

        return None

    # -----------------------------------------------------------------
    # Rendering
    # -----------------------------------------------------------------

    def _strip_front_matter(self, text: str) -> str:
        """Strip YAML front matter."""
        if text.startswith('---'):
            end = text.find('\n---', 3)
            if end != -1:
                return text[end + 4:].lstrip('\n')
        return text

    def _find_anchor(self, doc, anchor_heading: Optional[str]) -> int:
        """Find paragraph index after anchor heading. Returns 0 if not found."""
        if not anchor_heading:
            return len(doc.paragraphs)

        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == anchor_heading.strip():
                return i + 1
        return len(doc.paragraphs)

    def _render_markdown(
        self,
        doc,
        content: str,
        insert_idx: int,
        chapter_prefix: Optional[str],
        crossref_mode: str,
        first_line_indent: float,
        page_break_before_h2: bool,
        project_root: Path,
    ) -> None:
        """Convert markdown content to Word paragraphs and insert into doc."""
        from directives.parser import parse_directives, DirectiveType
        from directives.docx_helpers import (
            make_seq_field, make_ref_field, latex_to_omml,
            add_page_break, set_paragraph_style,
        )

        lines = content.split('\n')
        directives = parse_directives(content)
        directive_lines = {d.line for d in directives}

        # Build a map from line number to directive
        dir_map = {}
        for d in directives:
            dir_map[d.line] = d

        body = doc.element.body
        # Get reference element for insertion
        existing_paras = list(body.iterchildren(qn('w:p')))
        ref_element = existing_paras[insert_idx] if insert_idx < len(existing_paras) else None

        current_line = 0
        paragraph_buffer: List[str] = []

        def flush_paragraph():
            nonlocal paragraph_buffer
            if not paragraph_buffer:
                return
            text = '\n'.join(paragraph_buffer).strip()
            paragraph_buffer = []
            if not text:
                return

            p = doc.add_paragraph()
            # Apply first-line indent
            if first_line_indent > 0:
                pf = p.paragraph_format
                pf.first_line_indent = Pt(first_line_indent)
            self._add_inline_markup(p, text)
            self._move_paragraph(body, p._element, ref_element)

        def insert_paragraph(text: str = '', style: Optional[str] = None):
            p = doc.add_paragraph(text, style=style)
            self._move_paragraph(body, p._element, ref_element)
            return p

        i = 0
        while i < len(lines):
            line_num = i + 1  # 1-based
            line = lines[i]

            # Check if this line is part of a directive
            if line_num in dir_map:
                flush_paragraph()
                d = dir_map[line_num]

                if d.type == DirectiveType.PAGEBREAK:
                    p = doc.add_paragraph()
                    run = p.add_run()
                    br = etree.SubElement(run._element, qn('w:br'))
                    br.set(qn('w:type'), 'page')
                    self._move_paragraph(body, p._element, ref_element)

                elif d.type == DirectiveType.FIGURE:
                    self._insert_figure(
                        doc, body, ref_element, d,
                        chapter_prefix, crossref_mode, project_root,
                    )

                elif d.type == DirectiveType.EQUATION:
                    self._insert_equation(
                        doc, body, ref_element, d,
                        chapter_prefix, crossref_mode,
                    )

                elif d.type == DirectiveType.TABLE:
                    self._insert_table_caption(
                        doc, body, ref_element, d,
                        chapter_prefix, crossref_mode,
                    )
                    # Render the table body as markdown
                    if d.body:
                        self._render_table_body(doc, body, ref_element, d.body)

                elif d.type == DirectiveType.ALGORITHM:
                    self._insert_algorithm_caption(
                        doc, body, ref_element, d,
                        chapter_prefix, crossref_mode,
                    )
                    if d.body:
                        for bline in d.body.split('\n'):
                            insert_paragraph(bline, 'Normal')

                elif d.type == DirectiveType.REF:
                    label = d.label or ''
                    if crossref_mode == 'seq':
                        p = doc.add_paragraph()
                        for run_el in make_ref_field(f'_Ref_{label}'):
                            p._element.append(run_el)
                        self._move_paragraph(body, p._element, ref_element)
                    else:
                        ref_text = self._labels.get(label, f'[{label}]')
                        insert_paragraph(ref_text)

                elif d.type == DirectiveType.RAW_DOCX:
                    if d.body:
                        try:
                            raw_xml = etree.fromstring(d.body)
                            body.insert(
                                list(body).index(ref_element) if ref_element is not None else len(list(body)),
                                raw_xml,
                            )
                        except etree.XMLSyntaxError:
                            pass

                elif d.type == DirectiveType.STYLE:
                    p = insert_paragraph(d.style_text or '', d.style_name)

                # Skip lines consumed by block directive body
                if d.body is not None:
                    body_lines = d.body.count('\n') + 1
                    i += body_lines + 1  # +1 for end tag
                    i += 1
                    continue

                i += 1
                continue

            # Skip directive end tags
            stripped = line.strip()
            if re.match(r'^<!--\s*/(?:table|algorithm|raw-docx)\s*-->', stripped):
                i += 1
                continue

            # Headings
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                flush_paragraph()
                level = len(heading_match.group(1))
                heading_text = heading_match.group(2).strip()

                if page_break_before_h2 and level == 2:
                    p_br = doc.add_paragraph()
                    run = p_br.add_run()
                    br = etree.SubElement(run._element, qn('w:br'))
                    br.set(qn('w:type'), 'page')
                    self._move_paragraph(body, p_br._element, ref_element)

                p = insert_paragraph(heading_text, f'Heading {level}')
                i += 1
                continue

            # Blank line -> flush paragraph
            if not stripped:
                flush_paragraph()
                i += 1
                continue

            # Horizontal rule
            if re.match(r'^(-{3,}|\*{3,}|_{3,})$', stripped):
                flush_paragraph()
                i += 1
                continue

            # Accumulate paragraph text
            paragraph_buffer.append(line)
            i += 1

        flush_paragraph()

    # -----------------------------------------------------------------
    # Element insertion helpers
    # -----------------------------------------------------------------

    def _move_paragraph(self, body, p_element, ref_element):
        """Move a paragraph element to just before ref_element (or append)."""
        if ref_element is not None:
            body.insert(list(body).index(ref_element), p_element)
        else:
            body.append(p_element)

    def _add_inline_markup(self, paragraph, text: str) -> None:
        """Handle basic inline markdown: **bold**, *italic*, `code`."""
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                paragraph.add_run(part[2:-2]).bold = True
            elif part.startswith('*') and part.endswith('*'):
                paragraph.add_run(part[1:-1]).italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            else:
                paragraph.add_run(part)

    def _insert_figure(self, doc, body, ref_element, directive, prefix, crossref_mode, project_root):
        """Insert a figure with caption and auto-numbering."""
        from directives.docx_helpers import make_seq_field

        self._fig_counter += 1
        label = directive.label or ''
        img_path = directive.path or ''
        caption = directive.caption or ''
        width = directive.width

        # Resolve image path
        full_path = Path(img_path)
        if not full_path.is_absolute():
            full_path = project_root / img_path

        # Image paragraph
        if full_path.exists():
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_img.add_run()
            if width:
                try:
                    w = float(width.replace('%', '')) / 100 * 6  # ~6 inches content width
                    run.add_picture(str(full_path), width=Inches(w))
                except (ValueError, Exception):
                    run.add_picture(str(full_path))
            else:
                run.add_picture(str(full_path))
            self._move_paragraph(body, p_img._element, ref_element)

        # Caption paragraph
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

        num_text = f"{prefix}-{self._fig_counter}" if prefix else str(self._fig_counter)
        bookmark_name = f'_Ref_fig:{label}' if label else None
        self._labels[f'fig:{label}'] = f'\u56f3{num_text}'

        if crossref_mode == 'seq':
            p_cap.add_run('\u56f3')  # 図
            for el in make_seq_field('Figure', prefix, bookmark_name):
                p_cap._element.append(el)
            p_cap.add_run(f' {caption}')
        else:
            p_cap.add_run(f'\u56f3{num_text} {caption}')

        try:
            p_cap.style = 'Caption'
        except KeyError:
            pass

        self._move_paragraph(body, p_cap._element, ref_element)

    def _insert_equation(self, doc, body, ref_element, directive, prefix, crossref_mode):
        """Insert an equation with OMML and auto-numbering."""
        from directives.docx_helpers import make_seq_field, latex_to_omml

        self._eq_counter += 1
        label = directive.label or ''
        latex = directive.latex or ''

        num_text = f"({prefix}-{self._eq_counter})" if prefix else f"({self._eq_counter})"
        bookmark_name = f'_Ref_eq:{label}' if label else None
        self._labels[f'eq:{label}'] = num_text

        p = doc.add_paragraph()

        # Try OMML conversion
        omml = latex_to_omml(latex)
        if omml is not None:
            p._element.append(omml)
        else:
            # Fallback: just show LaTeX as text
            p.add_run(latex).italic = True

        # Equation number
        p.add_run('\t')
        if crossref_mode == 'seq':
            p.add_run('(')
            for el in make_seq_field('Equation', prefix, bookmark_name):
                p._element.append(el)
            p.add_run(')')
        else:
            p.add_run(num_text)

        self._move_paragraph(body, p._element, ref_element)

    def _insert_table_caption(self, doc, body, ref_element, directive, prefix, crossref_mode):
        """Insert a table caption with auto-numbering."""
        from directives.docx_helpers import make_seq_field

        self._tbl_counter += 1
        label = directive.label or ''
        caption = directive.caption or ''

        num_text = f"{prefix}-{self._tbl_counter}" if prefix else str(self._tbl_counter)
        bookmark_name = f'_Ref_tbl:{label}' if label else None
        self._labels[f'tbl:{label}'] = f'\u8868{num_text}'

        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if crossref_mode == 'seq':
            p_cap.add_run('\u8868')  # 表
            for el in make_seq_field('Table', prefix, bookmark_name):
                p_cap._element.append(el)
            p_cap.add_run(f' {caption}')
        else:
            p_cap.add_run(f'\u8868{num_text} {caption}')

        try:
            p_cap.style = 'Caption'
        except KeyError:
            pass

        self._move_paragraph(body, p_cap._element, ref_element)

    def _insert_algorithm_caption(self, doc, body, ref_element, directive, prefix, crossref_mode):
        """Insert an algorithm caption."""
        from directives.docx_helpers import make_seq_field

        self._alg_counter += 1
        label = directive.label or ''
        caption = directive.caption or ''

        num_text = f"{prefix}-{self._alg_counter}" if prefix else str(self._alg_counter)
        bookmark_name = f'_Ref_alg:{label}' if label else None
        self._labels[f'alg:{label}'] = f'Algorithm {num_text}'

        p_cap = doc.add_paragraph()
        p_cap.add_run(f'Algorithm {num_text}: {caption}').bold = True

        try:
            p_cap.style = 'Caption'
        except KeyError:
            pass

        self._move_paragraph(body, p_cap._element, ref_element)

    def _render_table_body(self, doc, body, ref_element, table_md: str) -> None:
        """Parse a simple markdown table and create a Word table."""
        lines = [l.strip() for l in table_md.strip().split('\n') if l.strip()]
        if len(lines) < 2:
            return

        # Parse header
        header_cells = [c.strip() for c in lines[0].strip('|').split('|')]
        num_cols = len(header_cells)

        # Skip separator line (line with dashes)
        data_start = 1
        if data_start < len(lines) and re.match(r'^[\|\s\-:]+$', lines[data_start]):
            data_start = 2

        # Parse data rows
        data_rows = []
        for row_line in lines[data_start:]:
            cells = [c.strip() for c in row_line.strip('|').split('|')]
            # Pad or trim to match header columns
            while len(cells) < num_cols:
                cells.append('')
            data_rows.append(cells[:num_cols])

        # Create Word table
        table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
        table.style = 'Table Grid'

        # Header row
        for j, text in enumerate(header_cells):
            cell = table.rows[0].cells[j]
            cell.text = text
            for run in cell.paragraphs[0].runs:
                run.bold = True

        # Data rows
        for i, row_data in enumerate(data_rows):
            for j, text in enumerate(row_data):
                table.rows[i + 1].cells[j].text = text

        # Move table element into position
        tbl_el = table._tbl
        if ref_element is not None:
            body.insert(list(body).index(ref_element), tbl_el)
        else:
            body.append(tbl_el)
